import streamlit as st
import pandas as pd
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import plotly.express as px
import plotly.graph_objects as go
import io
from datetime import datetime

# ================= PAGE CONFIG =================
st.set_page_config(
    page_title="Intelligent Data Analytics System",
    page_icon="🔬",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# ================= GLOBAL CSS =================
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Syne:wght@400;600;700;800&family=DM+Sans:wght@300;400;500;600&display=swap');

*, *::before, *::after { box-sizing: border-box; }

html, body, [data-testid="stAppViewContainer"] {
    background: #0a0a0f !important;
    color: #e8e8f0 !important;
    font-family: 'DM Sans', sans-serif !important;
}
[data-testid="stAppViewContainer"] {
    background: radial-gradient(ellipse 80% 60% at 50% -10%, #1a1040 0%, #0a0a0f 60%) !important;
}
#MainMenu, footer, header { visibility: hidden; }
[data-testid="stToolbar"] { display: none; }
.block-container { padding: 0 2.5rem 4rem !important; max-width: 1400px !important; }

::-webkit-scrollbar { width: 6px; height: 6px; }
::-webkit-scrollbar-track { background: #13131a; }
::-webkit-scrollbar-thumb { background: #3d3d6b; border-radius: 3px; }

/* HERO */
.hero-wrap { text-align: center; padding: 4rem 2rem 2.5rem; }
.hero-badge {
    display: inline-block;
    background: linear-gradient(135deg, rgba(99,110,250,0.15), rgba(178,75,243,0.15));
    border: 1px solid rgba(99,110,250,0.35);
    color: #a78bfa; font-family: 'DM Sans', sans-serif;
    font-size: 0.78rem; font-weight: 500; letter-spacing: 0.04em;
    padding: 0.4rem 1.2rem; border-radius: 999px; margin-bottom: 1.2rem;
}
.hero-title {
    font-family: 'Syne', sans-serif !important; font-size: clamp(2.4rem, 5vw, 3.8rem) !important;
    font-weight: 800 !important; line-height: 1.1 !important;
    background: linear-gradient(135deg, #ffffff 0%, #c4b5fd 50%, #818cf8 100%);
    -webkit-background-clip: text !important; -webkit-text-fill-color: transparent !important;
    background-clip: text !important; margin: 0 0 1rem !important; padding: 0 !important;
}
.hero-sub { color: #8b8ba7; font-size: 1.05rem; font-weight: 400; max-width: 560px; margin: 0 auto; line-height: 1.7; text-align: center !important; display: block; width: 100%; }
.hero-wrap p { text-align: center !important; }
.hero-wrap [data-testid="stMarkdownContainer"] p,
.hero-wrap [data-testid="stMarkdownContainer"],
[data-testid="stMarkdownContainer"] .hero-sub {
    text-align: center !important;
    width: 100% !important;
    display: block !important;
}
div:has(> .hero-wrap) p { text-align: center !important; }
.hero-divider { width: 80px; height: 3px; background: linear-gradient(90deg, #636efa, #b24bf3); border-radius: 2px; margin: 1.8rem auto 0; }

/* UPLOAD CARD */
.upload-card {
    background: linear-gradient(135deg, rgba(255,255,255,0.04), rgba(255,255,255,0.01));
    border: 1px solid rgba(99,110,250,0.25); border-radius: 20px; padding: 2rem 2rem 1rem;
    margin: 2rem auto 0.5rem; text-align: center;
    backdrop-filter: blur(12px); box-shadow: 0 0 60px rgba(99,110,250,0.08), inset 0 1px 0 rgba(255,255,255,0.06);
}
.upload-card-title { font-family: 'Syne', sans-serif; font-size: 1.1rem; font-weight: 700; color: #c4b5fd; margin-bottom: 0.4rem; }
.upload-card-sub { color: #5a5a7a; font-size: 0.82rem; margin-bottom: 0.8rem; }

/* SECTION HEADER */
.section-header {
    display: flex; align-items: center; gap: 0.75rem;
    margin: 2.8rem 0 1.2rem; padding-bottom: 0.75rem;
    border-bottom: 1px solid rgba(99,110,250,0.15);
}
.section-icon {
    width: 36px; height: 36px;
    background: linear-gradient(135deg, #636efa22, #b24bf322);
    border: 1px solid rgba(99,110,250,0.3); border-radius: 10px;
    display: flex; align-items: center; justify-content: center; font-size: 1rem; flex-shrink: 0;
}
.section-title { font-family: 'Syne', sans-serif !important; font-size: 1.25rem !important; font-weight: 700 !important; color: #e8e8f0 !important; margin: 0 !important; padding: 0 !important; }
.section-step { margin-left: auto; font-size: 0.72rem; font-weight: 600; letter-spacing: 0.1em; text-transform: uppercase; color: #4a4a6a; }

/* STAT CARDS */
.stat-row { display: grid; grid-template-columns: repeat(3, 1fr); gap: 1rem; margin: 1rem 0; }
.stat-card {
    background: rgba(255,255,255,0.03); border: 1px solid rgba(255,255,255,0.07);
    border-radius: 14px; padding: 1.25rem 1.4rem; position: relative; overflow: hidden;
}
.stat-card::before { content: ''; position: absolute; top: 0; left: 0; right: 0; height: 2px; background: var(--accent, linear-gradient(90deg,#636efa,#818cf8)); }
.stat-label { font-size: 0.72rem; font-weight: 600; letter-spacing: 0.1em; text-transform: uppercase; color: #5a5a7a; margin-bottom: 0.5rem; }
.stat-value { font-family: 'Syne', sans-serif; font-size: 2rem; font-weight: 800; color: #ffffff; line-height: 1; }
.stat-unit { font-size: 0.9rem; color: #6b6b8a; font-weight: 400; }

/* HEALTH BAR */
.health-bar-wrap { background: rgba(255,255,255,0.03); border: 1px solid rgba(255,255,255,0.07); border-radius: 14px; padding: 1.5rem 1.8rem; margin: 1rem 0; }
.health-bar-label { display: flex; justify-content: space-between; margin-bottom: 0.6rem; }
.health-bar-track { background: rgba(255,255,255,0.06); border-radius: 999px; height: 8px; overflow: hidden; margin-bottom: 0.75rem; }
.health-bar-fill { height: 100%; border-radius: 999px; transition: width 0.6s ease; }

/* BOXES */
.box { border-radius: 12px; padding: 0.9rem 1.2rem; margin: 0.6rem 0; font-size: 0.9rem; line-height: 1.5; display: flex; align-items: flex-start; gap: 0.6rem; }
.box-success { background: rgba(34,197,94,0.08); border: 1px solid rgba(34,197,94,0.2); color: #86efac; }
.box-warning { background: rgba(234,179,8,0.08); border: 1px solid rgba(234,179,8,0.2); color: #fde047; }
.box-info { background: rgba(99,110,250,0.08); border: 1px solid rgba(99,110,250,0.2); color: #a5b4fc; }
.box-error { background: rgba(239,68,68,0.08); border: 1px solid rgba(239,68,68,0.2); color: #fca5a5; }

/* INSIGHT CARD */
.insight-card { background: linear-gradient(135deg, rgba(99,110,250,0.07), rgba(178,75,243,0.05)); border: 1px solid rgba(99,110,250,0.2); border-radius: 16px; padding: 1.5rem 1.8rem; margin: 1rem 0; }
.insight-card-title { font-family: 'Syne', sans-serif; font-size: 0.72rem; font-weight: 700; letter-spacing: 0.12em; text-transform: uppercase; color: #818cf8; margin-bottom: 0.6rem; }
.insight-card-body { font-size: 0.97rem; color: #c4c4dc; line-height: 1.65; }

/* PILL */
.pill { display: inline-block; background: rgba(99,110,250,0.12); border: 1px solid rgba(99,110,250,0.25); color: #818cf8; font-size: 0.75rem; font-weight: 600; padding: 0.2rem 0.75rem; border-radius: 999px; margin-bottom: 0.6rem; }

/* PREVIEW BANNER */
.preview-banner { background: linear-gradient(135deg, rgba(99,110,250,0.1), rgba(178,75,243,0.08)); border: 1px solid rgba(99,110,250,0.25); border-radius: 16px; padding: 1.2rem 1.8rem; margin: 1.5rem 0; display: flex; align-items: center; justify-content: space-between; flex-wrap: wrap; gap: 0.5rem; }
.preview-banner-title { font-family: 'Syne', sans-serif; font-size: 1.1rem; font-weight: 700; color: #c4b5fd; }
.preview-banner-date { font-size: 0.82rem; color: #6b6b8a; }

/* BUTTONS */
.stDownloadButton > button {
    width: 100% !important; background: linear-gradient(135deg, #636efa, #818cf8) !important;
    color: white !important; border: none !important; border-radius: 12px !important;
    padding: 0.8rem 1.5rem !important; font-family: 'DM Sans', sans-serif !important;
    font-weight: 600 !important; font-size: 0.95rem !important;
    box-shadow: 0 4px 20px rgba(99,110,250,0.3) !important; transition: all 0.2s !important;
}
.stDownloadButton > button:hover { transform: translateY(-1px) !important; box-shadow: 0 6px 28px rgba(99,110,250,0.45) !important; }

.stButton > button[kind="primary"] {
    background: linear-gradient(135deg, #636efa, #b24bf3) !important; color: white !important;
    border: none !important; border-radius: 12px !important; padding: 0.75rem 2rem !important;
    font-family: 'DM Sans', sans-serif !important; font-weight: 600 !important; font-size: 1rem !important;
    width: 100% !important; box-shadow: 0 4px 24px rgba(99,110,250,0.35) !important; transition: all 0.2s !important;
}
.stButton > button[kind="primary"]:hover { transform: translateY(-2px) !important; box-shadow: 0 8px 32px rgba(99,110,250,0.5) !important; }
.stButton > button:not([kind="primary"]) { background: rgba(239,68,68,0.1) !important; color: #fca5a5 !important; border: 1px solid rgba(239,68,68,0.25) !important; border-radius: 10px !important; font-family: 'DM Sans', sans-serif !important; }

/* SELECTBOX */
[data-testid="stSelectbox"] > div > div { background: rgba(255,255,255,0.04) !important; border: 1px solid rgba(99,110,250,0.25) !important; border-radius: 10px !important; color: #e8e8f0 !important; }

/* DATAFRAME */
[data-testid="stDataFrame"] { border-radius: 14px !important; overflow: hidden !important; border: 1px solid rgba(255,255,255,0.07) !important; }

/* FILE UPLOADER */
[data-testid="stFileUploader"] * { text-align: center !important; }
[data-testid="stFileUploader"] > div { display: flex !important; flex-direction: column !important; align-items: center !important; justify-content: center !important; }
[data-testid="stFileUploader"] label { display: none !important; }
section[data-testid="stFileUploadDropzone"] {
    background: rgba(99,110,250,0.06) !important;
    border: 1.5px dashed rgba(99,110,250,0.35) !important;
    border-radius: 12px !important;
    display: flex !important;
    flex-direction: row !important;
    align-items: center !important;
    justify-content: center !important;
    gap: 1rem !important;
    padding: 1.2rem 2rem !important;
    width: 100% !important;
}
section[data-testid="stFileUploadDropzone"]:hover { border-color: rgba(99,110,250,0.65) !important; }
section[data-testid="stFileUploadDropzone"] button {
    background: linear-gradient(135deg, #636efa, #818cf8) !important;
    color: white !important; border: none !important; border-radius: 10px !important;
    padding: 0.55rem 1.5rem !important; font-family: 'DM Sans', sans-serif !important;
    font-weight: 600 !important; font-size: 0.9rem !important; flex-shrink: 0 !important; cursor: pointer !important;
}
section[data-testid="stFileUploadDropzone"] span,
section[data-testid="stFileUploadDropzone"] small,
section[data-testid="stFileUploadDropzone"] p {
    color: #5a5a7a !important; font-size: 0.82rem !important; text-align: left !important; margin: 0 !important; flex-shrink: 0 !important;
}
[data-testid="uploadedFileData"] {
    background: rgba(99,110,250,0.08) !important; border: 1px solid rgba(99,110,250,0.25) !important; border-radius: 10px !important; margin-top: 0.5rem !important;
}

hr { border-color: rgba(255,255,255,0.06) !important; }
[data-testid="stCaptionContainer"] { color: #4a4a6a !important; font-size: 0.8rem !important; }
</style>
""", unsafe_allow_html=True)


# ================= HERO =================
st.markdown("""
<div class="hero-wrap">
    <div class="hero-title">Intelligent Data<br>Analytics System</div>
    <p style="color:#8b8ba7;font-size:1.05rem;font-weight:400;line-height:1.7;text-align:center;width:100%;display:block;margin:0 auto;padding:0;">Upload your dataset — get instant cleaning, deep analysis,<br>beautiful visualizations, and exportable reports.</p>
    <div class="hero-divider"></div>
</div>

<a href="https://github.com/maazsaleem123/Intelligent-Data-Analytics-System" target="_blank"
   style="position:fixed;top:14px;left:16px;z-index:9999;display:inline-flex;align-items:center;gap:0.4rem;
          padding:0.35rem 0.85rem;background:rgba(20,20,30,0.85);border:1px solid rgba(99,110,250,0.3);
          border-radius:8px;color:#8b8ba7;font-family:'DM Sans',sans-serif;font-size:0.78rem;font-weight:500;
          text-decoration:none;backdrop-filter:blur(8px);"
   onmouseover="this.style.borderColor='rgba(99,110,250,0.7)';this.style.color='#c4b5fd';"
   onmouseout="this.style.borderColor='rgba(99,110,250,0.3)';this.style.color='#8b8ba7';">
    <svg height="15" width="15" viewBox="0 0 16 16" fill="currentColor" xmlns="http://www.w3.org/2000/svg">
      <path d="M8 0C3.58 0 0 3.58 0 8c0 3.54 2.29 6.53 5.47 7.59.4.07.55-.17.55-.38 0-.19-.01-.82-.01-1.49-2.01.37-2.53-.49-2.69-.94-.09-.23-.48-.94-.82-1.13-.28-.15-.68-.52-.01-.53.63-.01 1.08.58 1.23.82.72 1.21 1.87.87 2.33.66.07-.52.28-.87.51-1.07-1.78-.2-3.64-.89-3.64-3.95 0-.87.31-1.59.82-2.15-.08-.2-.36-1.02.08-2.12 0 0 .67-.21 2.2.82.64-.18 1.32-.27 2-.27.68 0 1.36.09 2 .27 1.53-1.04 2.2-.82 2.2-.82.44 1.1.16 1.92.08 2.12.51.56.82 1.27.82 2.15 0 3.07-1.87 3.75-3.65 3.95.29.25.54.73.54 1.48 0 1.07-.01 1.93-.01 2.2 0 .21.15.46.55.38A8.013 8.013 0 0016 8c0-4.42-3.58-8-8-8z"/>
    </svg>
    Source Code
</a>
""", unsafe_allow_html=True)


# ================= UPLOAD =================
col1, col2, col3 = st.columns([1, 2, 1])
with col2:
    st.markdown("""
    <div class="upload-card">
        <div class="upload-card-title">📂 Upload Your Dataset</div>
        <div class="upload-card-sub">Click below or drag & drop &nbsp;·&nbsp; CSV &nbsp;·&nbsp; Excel (.xlsx)</div>
    </div>
    """, unsafe_allow_html=True)
    uploaded_file = st.file_uploader("", type=["csv", "xlsx"], label_visibility="collapsed")


# ================= HELPERS =================
def get_chart_image(data, selected_col, chart_type):
    chart_data = data[selected_col].value_counts().reset_index()
    chart_data.columns = [selected_col, "Count"]
    is_numeric = pd.api.types.is_numeric_dtype(data[selected_col])
    fig, ax = plt.subplots(figsize=(11, 5))
    fig.patch.set_facecolor('#0d0d16'); ax.set_facecolor('#0d0d16')
    c = '#636EFA'
    if chart_type == "Bar Chart":
        ax.bar(chart_data[selected_col].astype(str), chart_data["Count"], color=c, edgecolor='#1a1a2e', linewidth=0.5)
        ax.set_xlabel(selected_col, fontsize=11, labelpad=10, color='#8b8ba7')
        ax.set_ylabel("Number of Records", fontsize=11, labelpad=10, color='#8b8ba7')
        ax.set_title(f"{selected_col} — Value Distribution", fontsize=13, pad=15, color='#e8e8f0')
        plt.xticks(rotation=30, ha='right', fontsize=9, color='#6b6b8a'); plt.yticks(fontsize=9, color='#6b6b8a')
        ax.grid(axis='y', linestyle='--', alpha=0.12, color='white'); ax.spines[['top','right','left','bottom']].set_color('#2a2a3e')
    elif chart_type == "Pie Chart":
        colors_p = ['#636EFA','#EF553B','#00CC96','#AB63FA','#FFA15A','#19D3F3','#FF6692']
        ax.pie(chart_data["Count"], labels=chart_data[selected_col].astype(str), autopct='%1.1f%%', startangle=140,
               textprops={'fontsize':9,'color':'#c4c4dc'}, colors=colors_p[:len(chart_data)])
        ax.set_title(f"{selected_col} — Percentage Distribution", fontsize=13, pad=15, color='#e8e8f0')
    elif chart_type == "Line Chart":
        ax.plot(chart_data[selected_col].astype(str), chart_data["Count"], marker='o', color=c, linewidth=2.5, markersize=7)
        ax.fill_between(range(len(chart_data)), chart_data["Count"], alpha=0.08, color=c)
        ax.set_xlabel(selected_col, fontsize=11, labelpad=10, color='#8b8ba7')
        ax.set_ylabel("Number of Records", fontsize=11, labelpad=10, color='#8b8ba7')
        ax.set_title(f"{selected_col} — Trend View", fontsize=13, pad=15, color='#e8e8f0')
        plt.xticks(rotation=30, ha='right', fontsize=9, color='#6b6b8a'); plt.yticks(fontsize=9, color='#6b6b8a')
        ax.grid(linestyle='--', alpha=0.12, color='white'); ax.spines[['top','right','left','bottom']].set_color('#2a2a3e')
    elif chart_type == "Area Chart":
        ax.fill_between(range(len(chart_data)), chart_data["Count"], alpha=0.25, color=c)
        ax.plot(range(len(chart_data)), chart_data["Count"], color=c, linewidth=2.5)
        ax.set_xticks(range(len(chart_data))); ax.set_xticklabels(chart_data[selected_col].astype(str), rotation=30, ha='right', fontsize=9, color='#6b6b8a')
        ax.set_xlabel(selected_col, fontsize=11, labelpad=10, color='#8b8ba7')
        ax.set_ylabel("Number of Records", fontsize=11, labelpad=10, color='#8b8ba7')
        ax.set_title(f"{selected_col} — Area View", fontsize=13, pad=15, color='#e8e8f0')
        ax.grid(axis='y', linestyle='--', alpha=0.12, color='white'); ax.spines[['top','right','left','bottom']].set_color('#2a2a3e')
    elif chart_type == "Histogram" and is_numeric:
        ax.hist(data[selected_col], bins=10, color=c, edgecolor='#1a1a2e', linewidth=0.5)
        ax.set_xlabel(f"{selected_col} Values", fontsize=11, labelpad=10, color='#8b8ba7')
        ax.set_ylabel("Number of Records", fontsize=11, labelpad=10, color='#8b8ba7')
        ax.set_title(f"{selected_col} — Distribution", fontsize=13, pad=15, color='#e8e8f0')
        plt.xticks(fontsize=9, color='#6b6b8a'); plt.yticks(fontsize=9, color='#6b6b8a')
        ax.grid(axis='y', linestyle='--', alpha=0.12, color='white'); ax.spines[['top','right','left','bottom']].set_color('#2a2a3e')
    plt.tight_layout(pad=2.0)
    buf = io.BytesIO()
    plt.savefig(buf, format='png', dpi=150, bbox_inches='tight', facecolor='#0d0d16')
    buf.seek(0); plt.close()
    return buf.read()


def section_header(icon, title, step=None):
    step_html = f'<span class="section-step">Step {step}</span>' if step else ''
    st.markdown(f'<div class="section-header"><div class="section-icon">{icon}</div><span class="section-title">{title}</span>{step_html}</div>', unsafe_allow_html=True)


def box(text, kind="info"):
    icons = {"success":"✓","warning":"⚠","info":"ℹ","error":"✕"}
    st.markdown(f'<div class="box box-{kind}"><span>{icons[kind]}</span><span>{text}</span></div>', unsafe_allow_html=True)


PLOTLY_LAYOUT = dict(
    paper_bgcolor='rgba(0,0,0,0)', plot_bgcolor='rgba(0,0,0,0)',
    font=dict(family='DM Sans', color='#8b8ba7', size=12),
    title_font=dict(family='Syne', color='#e8e8f0', size=15),
    xaxis=dict(gridcolor='rgba(255,255,255,0.05)', linecolor='rgba(255,255,255,0.08)', tickcolor='rgba(255,255,255,0.08)', tickfont=dict(color='#6b6b8a')),
    yaxis=dict(gridcolor='rgba(255,255,255,0.05)', linecolor='rgba(255,255,255,0.08)', tickcolor='rgba(255,255,255,0.08)', tickfont=dict(color='#6b6b8a')),
    margin=dict(l=60, r=40, t=70, b=70),
)


def render_chart(chart_data, data, selected_col, chart_type, key_suffix=""):
    if chart_type == "Bar Chart":
        fig = px.bar(chart_data, x=selected_col, y="Count", title=f"{selected_col} — Value Distribution",
                     labels={selected_col: selected_col, "Count": "Number of Records"}, color_discrete_sequence=["#636EFA"])
        fig.update_layout(**PLOTLY_LAYOUT, height=480, xaxis_tickangle=-30)
        fig.update_traces(marker_line_color='rgba(0,0,0,0)', opacity=0.9)
        st.plotly_chart(fig, use_container_width=True, key=f"bar{key_suffix}")
    elif chart_type == "Line Chart":
        fig = px.line(chart_data, x=selected_col, y="Count", title=f"{selected_col} — Trend View",
                      labels={selected_col: selected_col, "Count": "Number of Records"}, markers=True)
        fig.update_layout(**PLOTLY_LAYOUT, height=480)
        fig.update_traces(line_color='#636EFA', marker_color='#818cf8', line_width=2.5)
        st.plotly_chart(fig, use_container_width=True, key=f"line{key_suffix}")
    elif chart_type == "Area Chart":
        fig = px.area(chart_data, x=selected_col, y="Count", title=f"{selected_col} — Area View",
                      labels={selected_col: selected_col, "Count": "Number of Records"})
        fig.update_layout(**PLOTLY_LAYOUT, height=480)
        fig.update_traces(fillcolor='rgba(99,110,250,0.15)', line_color='#636EFA', line_width=2.5)
        st.plotly_chart(fig, use_container_width=True, key=f"area{key_suffix}")
    elif chart_type == "Histogram":
        if pd.api.types.is_numeric_dtype(data[selected_col]):
            fig = go.Figure(go.Histogram(x=data[selected_col], nbinsx=10, marker_color='#636EFA', marker_line_color='rgba(0,0,0,0)'))
            fig.update_layout(**PLOTLY_LAYOUT, height=480, title=f"{selected_col} — Distribution",
                              xaxis_title=f"{selected_col} Values", yaxis_title="Number of Records")
            st.plotly_chart(fig, use_container_width=True, key=f"hist{key_suffix}")
        else:
            box("Histogram requires numeric data. Please select a numeric column.", "error")
    elif chart_type == "Pie Chart":
        fig = px.pie(chart_data, names=selected_col, values="Count", title=f"{selected_col} — Percentage Distribution",
                     color_discrete_sequence=px.colors.sequential.Plasma_r)
        fig.update_layout(**{**PLOTLY_LAYOUT, 'xaxis':{}, 'yaxis':{}}, height=500,
                          legend=dict(orientation="v", x=1.02, y=0.5, font=dict(color='#8b8ba7')))
        fig.update_traces(textposition='inside', textinfo='percent+label', insidetextorientation='radial', textfont=dict(color='white', size=11))
        st.plotly_chart(fig, use_container_width=True, key=f"pie{key_suffix}")


# ================= PDF REPORT =================
def generate_pdf_report(data, selected_col, chart_type, clean_percent, missing_percent, duplicate_percent,
                         value_counts, percent_df, top_percent, insight_text, key_obs_text,
                         smart_suggestion_text, chart_img_bytes, missing_df, column):
    try:
        from reportlab.lib.pagesizes import A4, landscape
        from reportlab.lib import colors
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, Table,
                                         TableStyle, Image as RLImage, HRFlowable)
        from reportlab.lib.enums import TA_CENTER, TA_LEFT
    except ImportError:
        return None

    PAGE     = landscape(A4)
    L_MARGIN = 0.6 * inch
    R_MARGIN = 0.6 * inch
    T_MARGIN = 0.6 * inch
    B_MARGIN = 0.6 * inch
    USABLE_W = PAGE[0] - L_MARGIN - R_MARGIN

    C_NAVY   = colors.HexColor('#1A1A5E')
    C_BODY   = colors.HexColor('#1A1A1A')
    C_GRAY   = colors.HexColor('#555555')
    C_GREEN_T= colors.HexColor('#146C43')
    C_GREEN_B= colors.HexColor('#D4EDDA')
    C_BLUE_T = colors.HexColor('#084C61')
    C_BLUE_B = colors.HexColor('#D1ECF1')
    C_RED_T  = colors.HexColor('#9B1C1C')
    C_RED_B  = colors.HexColor('#FDE8E8')
    C_TH_BG  = colors.HexColor('#1E3A8A')
    C_ALT    = colors.HexColor('#EEF2FF')
    C_WHITE  = colors.white
    C_DIVIDER= colors.HexColor('#CCCCCC')
    C_ACCENT = colors.HexColor('#636EFA')
    C_PILL_BG= colors.HexColor('#1E1E3F')
    C_PILL_T = colors.HexColor('#818cf8')

    FONT_REG  = 'Times-Roman'
    FONT_BOLD = 'Times-Bold'
    FONT_ITAL = 'Times-Italic'

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=PAGE,
                            leftMargin=L_MARGIN, rightMargin=R_MARGIN,
                            topMargin=T_MARGIN,  bottomMargin=B_MARGIN)

    def ps(name, font=FONT_REG, size=10, color=None, align=TA_LEFT,
           bold=False, italic=False, space_before=0, space_after=4,
           bg=None, leading=None, left_indent=0):
        from reportlab.lib.styles import ParagraphStyle as PS
        return PS(name,
            fontName  = FONT_BOLD if bold else FONT_ITAL if italic else font,
            fontSize  = size,
            textColor = color or C_BODY,
            alignment = align,
            spaceBefore = space_before,
            spaceAfter  = space_after,
            backColor   = bg,
            leading     = leading or (size * 1.4),
            leftIndent  = left_indent,
            rightIndent = left_indent,
        )

    S_TITLE   = ps('title',  font=FONT_BOLD, size=22, color=C_NAVY, align=TA_CENTER, space_after=4)
    S_DATE    = ps('date',   font=FONT_ITAL, size=10, color=C_GRAY, align=TA_CENTER, space_after=6)
    S_HEAD    = ps('head',   font=FONT_BOLD, size=13, color=C_NAVY, space_before=14, space_after=6)
    S_BODY    = ps('body',   size=10, color=C_BODY, space_after=3)
    S_SUCCESS = ps('succ',   font=FONT_BOLD, size=10, color=C_GREEN_T, bg=C_GREEN_B, space_after=4, left_indent=8)
    S_INFO    = ps('info',   font=FONT_BOLD, size=10, color=C_BLUE_T,  bg=C_BLUE_B,  space_after=4, left_indent=8)
    S_WARN    = ps('warn',   font=FONT_BOLD, size=10, color=C_RED_T,   bg=C_RED_B,   space_after=4, left_indent=8)
    S_PILL    = ps('pill',   font=FONT_BOLD, size=9,  color=C_PILL_T,  bg=C_PILL_BG, space_after=4, left_indent=4)
    S_FOOT    = ps('foot',   font=FONT_ITAL, size=9,  color=C_GRAY,   align=TA_CENTER, space_after=0)

    def mk_table(rows, col_widths=None):
        n = len(rows[0])
        if col_widths is None:
            col_widths = [USABLE_W / n] * n
        t = Table(rows, colWidths=col_widths, repeatRows=1)
        t.setStyle(TableStyle([
            ('BACKGROUND',   (0,0), (-1,0),  C_TH_BG),
            ('TEXTCOLOR',    (0,0), (-1,0),  C_WHITE),
            ('FONTNAME',     (0,0), (-1,-1), FONT_REG),
            ('FONTNAME',     (0,0), (-1,0),  FONT_BOLD),
            ('FONTSIZE',     (0,0), (-1,0),  8),
            ('FONTSIZE',     (0,1), (-1,-1), 8),
            ('ALIGN',        (0,0), (-1,-1), 'CENTER'),
            ('VALIGN',       (0,0), (-1,-1), 'MIDDLE'),
            ('TOPPADDING',   (0,0), (-1,-1), 3),
            ('BOTTOMPADDING',(0,0), (-1,-1), 3),
            ('GRID',         (0,0), (-1,-1), 0.4, C_DIVIDER),
            ('ROWBACKGROUNDS',(0,1),(-1,-1), [C_WHITE, C_ALT]),
        ]))
        return t

    def pill_label(text):
        return Paragraph(f"  {text}", S_PILL)

    def section_heading(number, title):
        label = f"{number}.  {title}" if number else title
        return [
            Spacer(1, 0.05*inch),
            Paragraph(f'<font color="#1A1A5E"><b>{label}</b></font>', S_HEAD),
            HRFlowable(width="100%", thickness=1.5, color=C_NAVY, spaceAfter=6),
        ]

    def status_box(text, kind='success'):
        icon  = '✔' if kind=='success' else '⚠' if kind=='warning' else 'ℹ'
        style = S_SUCCESS if kind=='success' else S_WARN if kind=='warning' else S_INFO
        return Paragraph(f"  {icon}  {text}", style)

    el = []

    # ── TITLE ──────────────────────────────────────────────────────────────
    el += [Spacer(1, 0.15*inch),
           Paragraph("Intelligent Data Analytics System", S_TITLE),
           Spacer(1, 0.05*inch),
           Paragraph(f"Report Generated:  {datetime.now().strftime('%d %B %Y,  %I:%M %p')}", S_DATE),
           HRFlowable(width="100%", thickness=2, color=C_ACCENT, spaceAfter=10),
           Spacer(1, 0.1*inch)]

    # ── 1. DATASET PREVIEW ─────────────────────────────────────────────────
    el += section_heading(1, "Dataset Preview")
    el += [Paragraph(f'<b>Total Rows:</b>  {data.shape[0]}    <b>Total Columns:</b>  {data.shape[1]}', S_BODY),
           Spacer(1, 0.08*inch)]
    n_c    = len(data.columns) + 1
    id_w   = 0.3 * inch
    rest_w = (USABLE_W - id_w) / len(data.columns)
    col_w  = [id_w] + [rest_w] * len(data.columns)
    t_data = [['#'] + list(data.columns)]
    for i, row in enumerate(data.itertuples(index=False), 1):
        t_data.append([str(i)] + [str(v) for v in row])
    t = Table(t_data, colWidths=col_w, repeatRows=1)
    t.setStyle(TableStyle([
        ('BACKGROUND',   (0,0), (-1,0),  C_TH_BG),('TEXTCOLOR', (0,0), (-1,0),  C_WHITE),
        ('FONTNAME',     (0,0), (-1,0),  FONT_BOLD),('FONTNAME', (0,1), (-1,-1), FONT_REG),
        ('FONTSIZE',     (0,0), (-1,0),  7),('FONTSIZE', (0,1), (-1,-1), 7),
        ('ALIGN',        (0,0), (-1,-1), 'CENTER'),('VALIGN',    (0,0), (-1,-1), 'MIDDLE'),
        ('TOPPADDING',   (0,0), (-1,-1), 4),('BOTTOMPADDING', (0,0), (-1,-1), 4),
        ('LEFTPADDING',  (0,0), (-1,-1), 4),('RIGHTPADDING',  (0,0), (-1,-1), 4),
        ('GRID',         (0,0), (-1,-1), 0.4, C_DIVIDER),
        ('ROWBACKGROUNDS',(0,1),(-1,-1), [C_WHITE, C_ALT]),
    ]))
    el += [t, Spacer(1, 0.06*inch)]
    el.append(status_box("Your original file is safe and unchanged. All cleaning and analysis is done on a separate working copy.", 'info'))
    el.append(Spacer(1, 0.1*inch))

    # ── 2. DATA CLEANING ───────────────────────────────────────────────────
    el += section_heading(2, "Data Cleaning")
    orig_missing   = data.isnull().sum().sum()
    orig_duplicates= data.duplicated().sum()
    if orig_missing > 0 or orig_duplicates > 0:
        el.append(status_box(f"Original uploaded data has — Missing Values: {orig_missing}  |  Duplicate Rows: {orig_duplicates}", 'warning'))
    else:
        el.append(status_box("Original uploaded data is perfectly clean — no missing values, no duplicates.", 'success'))

    el.append(pill_label("🔍  Missing Values"))
    mv_rows = [['Column Name', 'Missing Count']] + \
              [[str(r['Column Name']), str(r['Missing Count'])] for _, r in missing_df.iterrows()]
    el += [mk_table(mv_rows, [USABLE_W*0.7, USABLE_W*0.3]), Spacer(1, 0.05*inch)]
    if data.isnull().sum().sum() == 0:
        el.append(status_box("No missing values found — dataset is perfectly clean.", 'success'))
    else:
        el.append(status_box(f"Total missing values found: {data.isnull().sum().sum()} — needs removal for clean analysis.", 'warning'))

    el.append(pill_label("🔁  Duplicate Rows"))
    dup_count = data.duplicated().sum()
    if dup_count == 0:
        el.append(status_box("No duplicate rows — all records are unique.", 'success'))
    else:
        el.append(status_box(f"{dup_count} duplicate row(s) found.", 'warning'))
    el.append(Spacer(1, 0.1*inch))

    # ── 3. DATA HEALTH SCORE ────────────────────────────────────────────────
    el += section_heading(3, "Data Health Score")
    status_txt = "Excellent ✔" if clean_percent==100 else "Good" if clean_percent>=80 else "Needs Cleaning ⚠"
    clr_hex    = '#146C43' if clean_percent==100 else '#084C61' if clean_percent>=80 else '#9B1C1C'
    el.append(Paragraph(
        f'<b>Data Quality Score:</b> <font color="{clr_hex}"><b>{clean_percent}%</b></font>    '
        f'<b>Clean:</b> {clean_percent}%    '
        f'<b>Missing:</b> {missing_percent}%    '
        f'<b>Duplicates:</b> {duplicate_percent}%    '
        f'<b>Status:</b> <font color="{clr_hex}"><b>{status_txt}</b></font>', S_BODY))
    el.append(Spacer(1, 0.1*inch))

    # ── 4. DATA ANALYSIS ────────────────────────────────────────────────────
    el += section_heading(4, "Data Analysis")
    el.append(pill_label("Statistical Summary"))
    desc = data.describe().round(2)
    stat_w = 1.1 * inch
    val_w  = (USABLE_W - stat_w) / len(desc.columns)
    s_rows = [['Statistic'] + list(desc.columns)] + \
             [[str(idx)] + [str(v) for v in desc.loc[idx]] for idx in desc.index]
    el += [mk_table(s_rows, [stat_w] + [val_w]*len(desc.columns)), Spacer(1, 0.08*inch)]

    el.append(Paragraph(f'<b>Select column for analysis:</b>  {column}', S_BODY))
    el.append(Spacer(1, 0.04*inch))
    el.append(pill_label("Value Counts"))
    vc = data[column].value_counts().reset_index(); vc.columns = [column, 'Count']
    el += [mk_table([[column, 'Count']] + [[str(r[column]), str(r['Count'])] for _, r in vc.iterrows()],
                    [USABLE_W*0.75, USABLE_W*0.25]), Spacer(1, 0.1*inch)]

    # ── 5. VISUALIZATION ────────────────────────────────────────────────────
    el += section_heading(5, f"Visualization — {selected_col}  ({chart_type})")
    el.append(status_box("Graph shows actual counts of each value — exact data, no ranges.", 'info'))
    if chart_img_bytes:
        el += [RLImage(io.BytesIO(chart_img_bytes), width=USABLE_W*0.75, height=USABLE_W*0.38),
               Spacer(1, 0.08*inch)]

    el.append(pill_label("📋  Graph Data Table"))
    el.append(Paragraph("Exact count of each value in the selected column.", S_BODY))
    gd = data[selected_col].value_counts().reset_index(); gd.columns = [selected_col, 'Count']
    el += [mk_table([[selected_col, 'Count']] + [[str(r[selected_col]), str(r['Count'])] for _, r in gd.iterrows()],
                    [USABLE_W*0.75, USABLE_W*0.25]), Spacer(1, 0.1*inch)]

    # ── 6. ADVANCED RECOMMENDATIONS ─────────────────────────────────────────
    el += section_heading(6, "Advanced Recommendations")
    if len(value_counts) > 0:
        top_value = value_counts.idxmax()
        low_value = value_counts.idxmin()
        el.append(status_box(f"{top_value} is the highest performing value in '{selected_col}'.", 'success'))
        if len(value_counts) > 1:
            el.append(status_box(f"{low_value} is the lowest performing value — improvement needed.", 'warning'))
        el.append(Spacer(1, 0.04*inch))

        el.append(pill_label("🔝  Top 3 Values"))
        t3 = value_counts.head(3).reset_index(); t3.columns = [selected_col, 'Count']
        el += [mk_table([[selected_col,'Count']]+[[str(r[selected_col]),str(r['Count'])] for _,r in t3.iterrows()],
                        [USABLE_W*0.75, USABLE_W*0.25]), Spacer(1, 0.04*inch)]

        el.append(pill_label("🔻  Bottom 3 Values"))
        b3 = value_counts.tail(3).reset_index(); b3.columns = [selected_col, 'Count']
        el += [mk_table([[selected_col,'Count']]+[[str(r[selected_col]),str(r['Count'])] for _,r in b3.iterrows()],
                        [USABLE_W*0.75, USABLE_W*0.25]), Spacer(1, 0.04*inch)]

        el.append(pill_label("📊  Percentage Contribution"))
        el += [mk_table([['Value','Contribution (%)']]+[[str(r['Value']),str(r['Contribution (%)'])] for _,r in percent_df.iterrows()],
                        [USABLE_W*0.75, USABLE_W*0.25]), Spacer(1, 0.04*inch)]

        max_percent = (value_counts / value_counts.sum() * 100).max()
        if max_percent > 60:
            el.append(status_box(f"{top_value} dominates with {round(max_percent,2)}% — reduce dependency for better balance.", 'warning'))
        elif max_percent > 40:
            el.append(status_box(f"{top_value} has strong influence ({round(max_percent,2)}%) — maintain but diversify.", 'warning'))
        else:
            el.append(status_box("Data distribution is fairly balanced across all categories.", 'success'))

        el.append(Spacer(1, 0.04*inch))
        el.append(status_box(f"Smart Suggestion: {smart_suggestion_text}", 'info'))
    el.append(Spacer(1, 0.1*inch))

    # ── 7. AUTO INSIGHT SUMMARY ─────────────────────────────────────────────
    el += section_heading(7, "Auto Insight Summary")
    if insight_text:
        el.append(status_box(insight_text, 'success'))
    el.append(Spacer(1, 0.04*inch))
    el += section_heading("", "Key Observations")
    if key_obs_text:
        el.append(status_box(key_obs_text, 'info'))
    el.append(Spacer(1, 0.15*inch))

    # ── FOOTER ─────────────────────────────────────────────────────────────
    el += [HRFlowable(width="100%", thickness=0.8, color=C_DIVIDER, spaceAfter=6),
           Paragraph("Intelligent Data Analytics System", S_FOOT)]

    doc.build(el)
    return buf.getvalue()


# ================= WORD REPORT =================
def generate_word_report(data, selected_col, chart_type, clean_percent, missing_percent, duplicate_percent,
                          value_counts, percent_df, top_percent, insight_text, key_obs_text,
                          smart_suggestion_text, chart_img_bytes, missing_df, column):
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        return None

    doc = Document()
    for s in doc.sections:
        s.orientation = 1
        s.page_width   = Inches(11.69)
        s.page_height  = Inches(8.27)
        s.top_margin   = Inches(0.8)
        s.bottom_margin= Inches(0.8)
        s.left_margin  = Inches(0.9)
        s.right_margin = Inches(0.9)

    FONT      = 'Times New Roman'
    CLR_HEAD  = RGBColor(0x1A, 0x1A, 0x5E)
    CLR_BODY  = RGBColor(0x1A, 0x1A, 0x1A)
    CLR_GREEN = RGBColor(0x14, 0x6C, 0x43)
    CLR_BLUE  = RGBColor(0x08, 0x4C, 0x61)
    CLR_RED   = RGBColor(0x9B, 0x1C, 0x1C)
    CLR_TBHD  = RGBColor(0xFF, 0xFF, 0xFF)
    CLR_PILL  = RGBColor(0x63, 0x6E, 0xFA)
    FILL_HEAD = '1E3A8A'
    FILL_ALT  = 'EEF2FF'
    FILL_PILL = 'E8EAFF'

    def set_font(run, size=11, bold=False, color=None, italic=False):
        run.font.name  = FONT
        run.font.size  = Pt(size)
        run.bold       = bold
        run.italic     = italic
        if color: run.font.color.rgb = color

    def add_title(text, size=22):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text)
        set_font(run, size=size, bold=True, color=CLR_HEAD)

    def add_subtitle(text, size=11):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(text)
        set_font(run, size=size, color=RGBColor(0x55,0x55,0x55), italic=True)

    def add_section_heading(number, title):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after  = Pt(6)
        label = f"{number}.  {title}" if number else title
        run = p.add_run(label)
        set_font(run, size=13, bold=True, color=CLR_HEAD)
        pPr  = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bot  = OxmlElement('w:bottom')
        bot.set(qn('w:val'),   'single')
        bot.set(qn('w:sz'),    '6')
        bot.set(qn('w:space'), '1')
        bot.set(qn('w:color'), '1E3A8A')
        pBdr.append(bot); pPr.append(pBdr)

    def add_pill(text):
        """Render a pill label exactly like the app's .pill style."""
        p = doc.add_paragraph()
        p.paragraph_format.space_after  = Pt(4)
        p.paragraph_format.space_before = Pt(6)
        run = p.add_run(f"  {text}  ")
        set_font(run, size=9, bold=True, color=CLR_PILL)
        # Light blue-purple shading on the paragraph
        pPr  = p._p.get_or_add_pPr()
        shd  = OxmlElement('w:shd')
        shd.set(qn('w:fill'), FILL_PILL)
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:val'),  'clear')
        pPr.append(shd)

    def add_body(text, size=10, color=None):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(text)
        set_font(run, size=size, color=color or CLR_BODY)

    def add_label_value(label, value):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r1 = p.add_run(f"{label}:  "); set_font(r1, size=10, bold=True, color=CLR_HEAD)
        r2 = p.add_run(str(value));    set_font(r2, size=10, color=CLR_BODY)

    def add_status_box(text, status='success'):
        p = doc.add_paragraph()
        p.paragraph_format.space_after  = Pt(4)
        p.paragraph_format.left_indent  = Inches(0.2)
        icon  = '✔' if status=='success' else '⚠' if status=='warning' else 'ℹ'
        color = CLR_GREEN if status=='success' else CLR_RED if status=='warning' else CLR_BLUE
        run   = p.add_run(f"  {icon}  {text}")
        set_font(run, size=10, bold=True, color=color)

    def add_table(headers, rows, col_widths=None):
        n       = len(headers)
        usable  = Inches(9.89)
        if col_widths is None:
            col_widths = [usable / n] * n
        t = doc.add_table(rows=1, cols=n)
        t.style = 'Table Grid'
        hdr = t.rows[0].cells
        for i, h in enumerate(headers):
            hdr[i].width = col_widths[i]
            p = hdr[i].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.clear()
            run = p.add_run(str(h)); set_font(run, size=9, bold=True, color=CLR_TBHD)
            tc = hdr[i]._tc; tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), FILL_HEAD); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:val'), 'clear')
            tcPr.append(shd)
        for ri, row_data in enumerate(rows):
            rc   = t.add_row().cells
            fill = FILL_ALT if ri % 2 == 0 else 'FFFFFF'
            for i, v in enumerate(row_data):
                rc[i].width = col_widths[i]
                p = rc[i].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.clear()
                run = p.add_run(str(v)); set_font(run, size=9, color=CLR_BODY)
                tc = rc[i]._tc; tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:fill'), fill); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:val'), 'clear')
                tcPr.append(shd)
        doc.add_paragraph()

    def add_divider():
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        pPr  = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bot  = OxmlElement('w:bottom')
        bot.set(qn('w:val'),   'single')
        bot.set(qn('w:sz'),    '4')
        bot.set(qn('w:space'), '1')
        bot.set(qn('w:color'), 'CCCCCC')
        pBdr.append(bot); pPr.append(pBdr)

    # ── TITLE ───────────────────────────────────────────────────────────────
    doc.add_paragraph()
    add_title("Intelligent Data Analytics System", size=24)
    doc.add_paragraph()
    add_subtitle(f"Report Generated:  {datetime.now().strftime('%d %B %Y,  %I:%M %p')}", size=11)
    add_divider()
    doc.add_paragraph()

    # ── 1. DATASET PREVIEW ──────────────────────────────────────────────────
    add_section_heading(1, "Dataset Preview")
    add_label_value("Total Rows",    data.shape[0])
    add_label_value("Total Columns", data.shape[1])
    doc.add_paragraph()

    usable_w   = Inches(9.89)
    id_w       = Inches(0.35)
    rest_w     = (usable_w - id_w) / len(data.columns)
    col_widths = [id_w] + [rest_w] * len(data.columns)
    n_cols     = len(data.columns) + 1

    t = doc.add_table(rows=1, cols=n_cols)
    t.style = 'Table Grid'
    hdr = t.rows[0].cells
    for i, h in enumerate(['#'] + list(data.columns)):
        hdr[i].width = col_widths[i]
        p = hdr[i].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.clear()
        run = p.add_run(str(h)); set_font(run, size=8, bold=True, color=CLR_TBHD)
        tc = hdr[i]._tc; tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), FILL_HEAD); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:val'), 'clear')
        tcPr.append(shd)
    for idx, row in enumerate(data.itertuples(index=False), 1):
        rc   = t.add_row().cells
        fill = FILL_ALT if idx % 2 == 0 else 'FFFFFF'
        rc[0].width = col_widths[0]
        p = rc[0].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.clear()
        run = p.add_run(str(idx)); set_font(run, size=8, bold=True, color=CLR_HEAD)
        for i, v in enumerate(row, 1):
            rc[i].width = col_widths[i]
            p = rc[i].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.clear()
            run = p.add_run(str(v)); set_font(run, size=8, color=CLR_BODY)
            tc = rc[i]._tc; tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), fill); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:val'), 'clear')
            tcPr.append(shd)
    doc.add_paragraph()
    add_status_box("Your original file is safe and unchanged. All cleaning and analysis is done on a separate working copy.", 'info')

    # ── 2. DATA CLEANING ────────────────────────────────────────────────────
    add_section_heading(2, "Data Cleaning")
    orig_missing    = data.isnull().sum().sum()
    orig_duplicates = data.duplicated().sum()
    if orig_missing > 0 or orig_duplicates > 0:
        add_status_box(f"Original uploaded data has — Missing Values: {orig_missing}  |  Duplicate Rows: {orig_duplicates}", 'warning')
    else:
        add_status_box("Original uploaded data is perfectly clean — no missing values, no duplicates.", 'success')

    add_pill("🔍  Missing Values")
    mv_headers = ['Column Name', 'Missing Count']
    mv_rows    = [[str(r['Column Name']), str(r['Missing Count'])] for _, r in missing_df.iterrows()]
    add_table(mv_headers, mv_rows, [Inches(5.0), Inches(4.89)])
    if data.isnull().sum().sum() == 0:
        add_status_box("No missing values found — dataset is perfectly clean.", 'success')
    else:
        add_status_box(f"Total missing values found: {data.isnull().sum().sum()} — needs removal for clean analysis.", 'warning')

    add_pill("🔁  Duplicate Rows")
    dup_count = data.duplicated().sum()
    if dup_count == 0:
        add_status_box("No duplicate rows — all records are unique.", 'success')
    else:
        add_status_box(f"{dup_count} duplicate row(s) found.", 'warning')

    # ── 3. DATA HEALTH SCORE ────────────────────────────────────────────────
    add_section_heading(3, "Data Health Score")
    status = "Excellent ✔" if clean_percent==100 else "Good 👍" if clean_percent>=80 else "Needs Cleaning ⚠"
    color  = CLR_GREEN if clean_percent==100 else CLR_BLUE if clean_percent>=80 else CLR_RED
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    for label, val in [("Data Quality Score", f"{clean_percent}%"), ("  |  Clean", f"{clean_percent}%"),
                       ("  |  Missing", f"{missing_percent}%"), ("  |  Duplicates", f"{duplicate_percent}%"),
                       ("  |  Status", status)]:
        r1 = p.add_run(f"{label}: "); set_font(r1, size=11, bold=True, color=CLR_HEAD)
        r2 = p.add_run(val + "   "); set_font(r2, size=11, bold=True, color=color)

    # ── 4. DATA ANALYSIS ────────────────────────────────────────────────────
    add_section_heading(4, "Data Analysis")
    add_pill("Statistical Summary")
    desc = data.describe().round(2)
    s_headers  = ['Statistic'] + list(desc.columns)
    s_rows_data= [[str(idx)] + [str(v) for v in desc.loc[idx]] for idx in desc.index]
    stat_w     = Inches(1.2)
    val_w      = (Inches(9.89) - stat_w) / len(desc.columns)
    add_table(s_headers, s_rows_data, [stat_w] + [val_w]*len(desc.columns))

    add_body(f"Select column for analysis:  {column}", size=10, color=CLR_HEAD)
    add_pill("Value Counts")
    vc = data[column].value_counts().reset_index(); vc.columns = [column, 'Count']
    add_table([column, 'Count'],
              [[str(r[column]), str(r['Count'])] for _, r in vc.iterrows()],
              [Inches(6.0), Inches(3.89)])

    # ── 5. VISUALIZATION ────────────────────────────────────────────────────
    add_section_heading(5, f"Visualization — {selected_col}  ({chart_type})")
    add_status_box("Graph shows actual counts of each value — exact data, no ranges.", 'info')
    if chart_img_bytes:
        doc.add_picture(io.BytesIO(chart_img_bytes), width=Inches(7))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    add_pill("📋  Graph Data Table")
    add_body("Exact count of each value in the selected column.", size=10, color=CLR_BODY)
    gd = data[selected_col].value_counts().reset_index(); gd.columns = [selected_col, 'Count']
    add_table([selected_col, 'Count'],
              [[str(r[selected_col]), str(r['Count'])] for _, r in gd.iterrows()],
              [Inches(6.0), Inches(3.89)])

    # ── 6. ADVANCED RECOMMENDATIONS ─────────────────────────────────────────
    add_section_heading(6, "Advanced Recommendations")
    if len(value_counts) > 0:
        top_value = value_counts.idxmax()
        low_value = value_counts.idxmin()
        add_status_box(f"{top_value} is the highest performing value in '{selected_col}'.", 'success')
        if len(value_counts) > 1:
            add_status_box(f"{low_value} is the lowest performing value — improvement needed.", 'warning')
        doc.add_paragraph()

        add_pill("🔝  Top 3 Values")
        t3 = value_counts.head(3).reset_index(); t3.columns = [selected_col, 'Count']
        add_table([selected_col, 'Count'],
                  [[str(r[selected_col]), str(r['Count'])] for _, r in t3.iterrows()],
                  [Inches(6.0), Inches(3.89)])

        add_pill("🔻  Bottom 3 Values")
        b3 = value_counts.tail(3).reset_index(); b3.columns = [selected_col, 'Count']
        add_table([selected_col, 'Count'],
                  [[str(r[selected_col]), str(r['Count'])] for _, r in b3.iterrows()],
                  [Inches(6.0), Inches(3.89)])

        add_pill("📊  Percentage Contribution")
        add_table(['Value', 'Contribution (%)'],
                  [[str(r['Value']), str(r['Contribution (%)'])] for _, r in percent_df.iterrows()],
                  [Inches(6.0), Inches(3.89)])

        max_percent = (value_counts / value_counts.sum() * 100).max()
        if max_percent > 60:
            add_status_box(f"{top_value} dominates with {round(max_percent,2)}% — reduce dependency for better balance.", 'warning')
        elif max_percent > 40:
            add_status_box(f"{top_value} has strong influence ({round(max_percent,2)}%) — maintain but diversify.", 'warning')
        else:
            add_status_box("Data distribution is fairly balanced across all categories.", 'success')

        p = doc.add_paragraph()
        r1 = p.add_run("Smart Suggestion:  "); set_font(r1, size=10, bold=True, color=CLR_BLUE)
        r2 = p.add_run(smart_suggestion_text);  set_font(r2, size=10, color=CLR_BODY)

    # ── 7. AUTO INSIGHT SUMMARY ─────────────────────────────────────────────
    add_section_heading(7, "Auto Insight Summary")
    if insight_text:
        add_status_box(insight_text, 'success')
    add_section_heading("", "Key Observations")
    if key_obs_text:
        add_status_box(key_obs_text, 'info')

    # ── FOOTER ──────────────────────────────────────────────────────────────
    doc.add_paragraph()
    add_divider()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Intelligent Data Analytics System")
    set_font(run, size=9, italic=True, color=RGBColor(0x99,0x99,0x99))

    buf = io.BytesIO(); doc.save(buf); return buf.getvalue()


# ===================================================
# ================= MAIN LOGIC ======================
# ===================================================
if uploaded_file is not None:
    file_key = uploaded_file.name + str(uploaded_file.size)
    if 'file_key' not in st.session_state or st.session_state.file_key != file_key:
        if uploaded_file.name.endswith(".csv"):
            raw = pd.read_csv(uploaded_file)
        else:
            raw = pd.read_excel(uploaded_file)
        st.session_state.original_data = raw.copy()
        st.session_state.data = raw.copy()
        st.session_state.file_key = file_key

    data = st.session_state.data

    # ── 1. DATASET PREVIEW ──────────────────────────────────────────────────
    section_header("📄", "Dataset Preview", step=1)
    original_data = st.session_state.original_data.copy()
    original_data.index = range(1, len(original_data)+1)
    st.dataframe(original_data, use_container_width=True)
    st.markdown(f"""
    <div class="stat-row">
        <div class="stat-card" style="--accent:linear-gradient(90deg,#636efa,#818cf8)">
            <div class="stat-label">Total Rows</div>
            <div class="stat-value">{st.session_state.original_data.shape[0]}<span class="stat-unit"> rows</span></div>
        </div>
        <div class="stat-card" style="--accent:linear-gradient(90deg,#b24bf3,#636efa)">
            <div class="stat-label">Total Columns</div>
            <div class="stat-value">{st.session_state.original_data.shape[1]}<span class="stat-unit"> cols</span></div>
        </div>
        <div class="stat-card" style="--accent:linear-gradient(90deg,#22c55e,#16a34a)">
            <div class="stat-label">Status</div>
            <div class="stat-value" style="font-size:1.1rem;color:#86efac;padding-top:0.4rem">✓ Loaded</div>
        </div>
    </div>""", unsafe_allow_html=True)
    box("Your original file is safe and unchanged. All cleaning and analysis is done on a separate working copy.", "info")

    # ── 2. DATA CLEANING ────────────────────────────────────────────────────
    section_header("🧹", "Data Cleaning", step=2)
    orig            = st.session_state.original_data
    orig_missing    = orig.isnull().sum().sum()
    orig_duplicates = orig.duplicated().sum()

    if orig_missing > 0 or orig_duplicates > 0:
        box(f"Original uploaded data has — Missing Values: <b>{orig_missing}</b>  |  Duplicate Rows: <b>{orig_duplicates}</b>", "warning")
    else:
        box("Original uploaded data is perfectly clean — no missing values, no duplicates.", "success")

    st.markdown('<div class="pill">🔍 Missing Values</div>', unsafe_allow_html=True)
    missing    = data.isnull().sum()
    missing_df = missing.reset_index(); missing_df.columns=["Column Name","Missing Count"]; missing_df.index=range(1,len(missing_df)+1)
    st.dataframe(missing_df, use_container_width=True)

    if missing.sum()==0:
        box("✅ Missing values removed from working copy — analysis will use clean data.","success") if orig_missing > 0 else box("No missing values found — dataset is perfectly clean.","success")
    else:
        box(f"Total missing values found: {missing.sum()} — click below to remove for analysis.","warning")
        if st.button("🗑️ Remove Missing Values"):
            st.session_state.data = st.session_state.data.dropna()
            data = st.session_state.data
            st.rerun()

    st.markdown('<div class="pill" style="margin-top:1rem">🔁 Duplicate Rows</div>', unsafe_allow_html=True)
    dup_count = data.duplicated().sum()
    if dup_count==0:
        box("✅ Duplicate rows removed from working copy — analysis will use clean data.","success") if orig_duplicates > 0 else box("No duplicate rows — all records are unique.","success")
    else:
        box(f"{dup_count} duplicate row(s) found — click below to remove for analysis.","warning")
        if st.button("🗑️ Remove Duplicate Rows"):
            st.session_state.data = st.session_state.data.drop_duplicates()
            data = st.session_state.data
            st.rerun()

    if orig_missing > 0 or orig_duplicates > 0:
        st.markdown("<div style='margin-top:0.8rem'></div>", unsafe_allow_html=True)
        if st.button("🔄 Reset to Original Data"):
            st.session_state.data = st.session_state.original_data.copy()
            st.rerun()

    # ── 3. DATA HEALTH SCORE ────────────────────────────────────────────────
    section_header("💚", "Data Health Score", step=3)
    total_cells      = data.shape[0]*data.shape[1]
    missing_cells    = data.isnull().sum().sum()
    dup_rows         = data.duplicated().sum()
    missing_percent  = round((missing_cells/total_cells)*100,2)
    duplicate_percent= round((dup_rows/data.shape[0])*100,2)
    clean_percent    = round(max(100-missing_percent-duplicate_percent,0),2)
    health_color     = "#22c55e" if clean_percent==100 else "#eab308" if clean_percent>=80 else "#ef4444"
    health_label     = "Excellent 🏆" if clean_percent==100 else "Good 👍" if clean_percent>=80 else "Needs Cleaning ⚠️"
    st.markdown(f"""
    <div class="health-bar-wrap">
        <div class="health-bar-label">
            <span style="font-family:'Syne',sans-serif;font-weight:700;color:#e8e8f0">Data Quality Score</span>
            <span style="font-family:'Syne',sans-serif;font-size:1.4rem;font-weight:800;color:{health_color}">{clean_percent}%</span>
        </div>
        <div class="health-bar-track"><div class="health-bar-fill" style="width:{clean_percent}%;background:linear-gradient(90deg,#636efa,{health_color})"></div></div>
        <div style="display:flex;gap:2rem;font-size:0.82rem;color:#6b6b8a">
            <span>✓ Clean: <b style="color:#e8e8f0">{clean_percent}%</b></span>
            <span>⚠ Missing: <b style="color:#e8e8f0">{missing_percent}%</b></span>
            <span>⧉ Duplicates: <b style="color:#e8e8f0">{duplicate_percent}%</b></span>
            <span style="margin-left:auto;color:{health_color};font-weight:600">{health_label}</span>
        </div>
    </div>""", unsafe_allow_html=True)

    # ── 4. DATA ANALYSIS ────────────────────────────────────────────────────
    section_header("📊", "Data Analysis", step=4)
    st.markdown('<div class="pill">Statistical Summary</div>', unsafe_allow_html=True)
    st.dataframe(data.describe(), use_container_width=True)

    def is_id_column(col_name, series):
        name_lower = col_name.lower()
        id_kw = ["id","code","no","number","num","serial","ref","roll"]
        return any(kw in name_lower for kw in id_kw) or series.nunique()/len(series)==1.0

    column = st.selectbox("Select column for analysis", data.columns)
    if is_id_column(column, data[column]):
        box(f"'{column}' appears to be an ID/code column. Select a meaningful column like City, Product, Gender, etc.","warning")

    st.markdown('<div class="pill" style="margin-top:0.8rem">Value Counts</div>', unsafe_allow_html=True)
    vc_df = data[column].value_counts().reset_index(); vc_df.columns=[column,"Count"]; vc_df.index=range(1,len(vc_df)+1)
    st.dataframe(vc_df, use_container_width=True)

    # ── 5. VISUALIZATION ────────────────────────────────────────────────────
    section_header("📈", "Visualization", step=5)
    selected_col = st.selectbox("Select column for visualization", data.columns)
    if is_id_column(selected_col, data[selected_col]):
        box(f"'{selected_col}' looks like an ID/code column. Select a meaningful column.","warning")
    is_numeric = pd.api.types.is_numeric_dtype(data[selected_col])
    chart_type = st.selectbox("Select chart type", ["Bar Chart","Pie Chart","Line Chart","Area Chart","Histogram"])

    suggestion_text = ""
    if chart_type in ["Line Chart","Area Chart"] and not is_numeric:
        box("Line/Area Chart works better with numeric or time-based data.","warning"); suggestion_text="Suggested: Bar Chart or Pie Chart."
    elif chart_type=="Histogram" and not is_numeric:
        box("Histogram works best with numeric data.","warning"); suggestion_text="Suggested: Bar Chart or Pie Chart."
    elif chart_type=="Pie Chart" and is_numeric:
        box("Pie Chart is better suited for categorical data.","warning"); suggestion_text="Suggested: Histogram or Line Chart."
    if suggestion_text: box(f"💡 {suggestion_text}","info")

    box("Graph shows actual counts of each value — exact data, no ranges.","info")
    chart_data = data[selected_col].value_counts().reset_index(); chart_data.columns=[selected_col,"Count"]
    render_chart(chart_data, data, selected_col, chart_type, key_suffix="_main")

    st.markdown('<div class="pill" style="margin-top:0.5rem">📋 Graph Data Table</div>', unsafe_allow_html=True)
    st.caption("Exact count of each value in the selected column.")
    chart_data_display = chart_data.copy(); chart_data_display.index=range(1,len(chart_data_display)+1)
    st.dataframe(chart_data_display, use_container_width=True)

    # ── 6. ADVANCED RECOMMENDATIONS ─────────────────────────────────────────
    section_header("💡", "Advanced Recommendations", step=6)
    value_counts = data[selected_col].value_counts()
    total        = value_counts.sum()
    top_percent  = 0; insight_text=""; key_obs_text=""; smart_suggestion_text=""
    percent_df   = pd.DataFrame(columns=["Value","Contribution (%)"])

    if len(value_counts) > 0:
        top_value = value_counts.idxmax()
        low_value = value_counts.idxmin()
        box(f"<b>{top_value}</b> is the highest performing value in '{selected_col}'.","success")
        if len(value_counts)>1: box(f"<b>{low_value}</b> is the lowest performing value — improvement needed.","warning")

        col_a, col_b = st.columns(2)
        with col_a:
            st.markdown('<div class="pill">🔝 Top 3 Values</div>', unsafe_allow_html=True)
            top3 = value_counts.head(3).reset_index(); top3.columns=[selected_col,"Count"]; top3.index=range(1,len(top3)+1)
            st.dataframe(top3, use_container_width=True)
        with col_b:
            st.markdown('<div class="pill">🔻 Bottom 3 Values</div>', unsafe_allow_html=True)
            bottom3 = value_counts.tail(3).reset_index(); bottom3.columns=[selected_col,"Count"]; bottom3.index=range(1,len(bottom3)+1)
            st.dataframe(bottom3, use_container_width=True)

        st.markdown('<div class="pill" style="margin-top:1rem">📊 Percentage Contribution</div>', unsafe_allow_html=True)
        percentage = (value_counts/total*100).round(2)
        percent_df = percentage.reset_index(); percent_df.columns=["Value","Contribution (%)"]; percent_df.index=range(1,len(percent_df)+1)
        st.dataframe(percent_df, use_container_width=True)

        max_percent = percentage.max()
        if max_percent>60: box(f"<b>{top_value}</b> dominates with {max_percent}% — reduce dependency for better balance.","error")
        elif max_percent>40: box(f"<b>{top_value}</b> has strong influence ({max_percent}%) — maintain but diversify.","warning")
        else: box("Data distribution is fairly balanced across all categories.","success")

        col_lower = selected_col.lower()
        if "sales" in col_lower or "price" in col_lower or "revenue" in col_lower: smart_suggestion_text="Focus on high-value items to maximize revenue."
        elif "age" in col_lower: smart_suggestion_text="Target the most frequent age group for better reach."
        elif "quantity" in col_lower or "qty" in col_lower: smart_suggestion_text="Optimize inventory based on demand patterns."
        elif "city" in col_lower or "region" in col_lower or "location" in col_lower: smart_suggestion_text="Focus marketing efforts on top-performing locations."
        elif "product" in col_lower or "category" in col_lower: smart_suggestion_text="Scale top-performing products and improve weak categories."
        elif "gender" in col_lower: smart_suggestion_text="Review if offerings are balanced across genders."
        elif "payment" in col_lower: smart_suggestion_text="Promote the most used payment mode and consider removing less-used ones."
        elif "department" in col_lower: smart_suggestion_text="Invest in top-performing departments and support weaker ones."
        elif "rating" in col_lower or "score" in col_lower: smart_suggestion_text="Focus on improving lower ratings and maintaining top scores."
        elif "designation" in col_lower or "role" in col_lower: smart_suggestion_text="Review resource allocation across designations."
        else: smart_suggestion_text="Improve low-performing values and scale top-performing ones."
        st.markdown(f'<div class="insight-card"><div class="insight-card-title">🧠 Smart Suggestion</div><div class="insight-card-body">💡 {smart_suggestion_text}</div></div>', unsafe_allow_html=True)

    # ── 7. AUTO INSIGHT SUMMARY ──────────────────────────────────────────────
    section_header("🧠", "Auto Insight Summary", step=7)
    if len(value_counts) > 0:
        top_value    = value_counts.idxmax()
        top_count    = value_counts.max()
        total_records= value_counts.sum()
        top_percent  = round((top_count/total_records)*100,2)
        insight_text = f"'{top_value}' contributes {top_percent}% of the total data and is the top-performing category."
        if len(value_counts) > 1:
            second_value   = value_counts.index[1]
            second_percent = round((value_counts.iloc[1]/total_records)*100,2)
            insight_text  += f" It is followed by '{second_value}' with {second_percent}% contribution."
        insight_text += (" The dataset is highly concentrated in one category." if top_percent>60
                         else " The dataset shows moderate dominance by one category." if top_percent>40
                         else " The dataset is fairly balanced across categories.")
        st.markdown(f'<div class="insight-card"><div class="insight-card-title">📊 Insight</div><div class="insight-card-body">👉 {insight_text}</div></div>', unsafe_allow_html=True)

        col_lower = selected_col.lower()
        if "sales" in col_lower or "price" in col_lower:
            key_obs_text="Revenue is driven by high-value categories. Focus on scaling them."
        elif "age" in col_lower:
            key_obs_text="Most customers belong to a specific age group. Target marketing accordingly."
        elif "quantity" in col_lower:
            key_obs_text="Demand pattern suggests optimizing stock levels."
        elif "city" in col_lower or "region" in col_lower:
            key_obs_text="Certain locations show higher activity. Prioritize resources there."
        elif "product" in col_lower or "category" in col_lower:
            key_obs_text="Top product categories are clear. Invest more in them."
        else:
            key_obs_text="Focus on improving lower-performing categories for better overall balance."
        st.markdown(f'<div class="insight-card" style="margin-top:0.8rem;border-color:rgba(99,110,250,0.15)"><div class="insight-card-title">📌 Key Observations</div><div class="insight-card-body">💡 {key_obs_text}</div></div>', unsafe_allow_html=True)

    # ── 8. DOWNLOAD REPORT ────────────────────────────────────────────────────
    section_header("📥", "Download Report", step=8)
    st.markdown('<p style="color:#6b6b8a;font-size:0.9rem;margin-bottom:1.5rem">Preview your full analysis and download as Word or PDF.</p>', unsafe_allow_html=True)

    if st.button("👁️ Preview & Download Report", type="primary"):
        with st.spinner("Building your report..."):
            chart_img_bytes = get_chart_image(data, selected_col, chart_type)

        st.markdown(f"""
        <div class="preview-banner">
            <div><div class="preview-banner-title">📋 Report Preview</div><div class="preview-banner-date">Identical to your live analysis above.</div></div>
            <div class="preview-banner-date">Generated: {datetime.now().strftime('%d %B %Y, %I:%M %p')}</div>
        </div>""", unsafe_allow_html=True)

        # ── Preview: Step 1 ──────────────────────────────────────────────
        section_header("📄","Dataset Preview")
        prev = data.copy(); prev.index=range(1,len(prev)+1)
        st.dataframe(prev, use_container_width=True)
        box(f"Dataset Loaded: {data.shape[0]} rows, {data.shape[1]} columns","success")
        box("Your original file is safe and unchanged. All cleaning and analysis is done on a separate working copy.","info")

        # ── Preview: Step 2 ──────────────────────────────────────────────
        section_header("🧹","Data Cleaning")
        if orig_missing > 0 or orig_duplicates > 0:
            box(f"Original uploaded data has — Missing Values: <b>{orig_missing}</b>  |  Duplicate Rows: <b>{orig_duplicates}</b>","warning")
        else:
            box("Original uploaded data is perfectly clean — no missing values, no duplicates.","success")

        st.markdown('<div class="pill">🔍 Missing Values</div>', unsafe_allow_html=True)
        st.dataframe(missing_df, use_container_width=True)
        if data.isnull().sum().sum()==0:
            box("No missing values found — dataset is perfectly clean.","success")
        else:
            box(f"Total missing values found: {data.isnull().sum().sum()}","warning")

        st.markdown('<div class="pill" style="margin-top:0.8rem">🔁 Duplicate Rows</div>', unsafe_allow_html=True)
        dc = data.duplicated().sum()
        if dc==0: box("No duplicate rows — all records are unique.","success")
        else:     box(f"{dc} duplicate row(s) found.","warning")

        # ── Preview: Step 3 ──────────────────────────────────────────────
        section_header("💚","Data Health Score")
        st.markdown(f"""
        <div class="health-bar-wrap">
            <div class="health-bar-label">
                <span style="font-family:'Syne',sans-serif;font-weight:700;color:#e8e8f0">Data Quality Score</span>
                <span style="font-family:'Syne',sans-serif;font-size:1.4rem;font-weight:800;color:{health_color}">{clean_percent}%</span>
            </div>
            <div class="health-bar-track"><div class="health-bar-fill" style="width:{clean_percent}%;background:linear-gradient(90deg,#636efa,{health_color})"></div></div>
            <div style="display:flex;gap:2rem;font-size:0.82rem;color:#6b6b8a">
                <span>✓ Clean: <b style="color:#e8e8f0">{clean_percent}%</b></span>
                <span>⚠ Missing: <b style="color:#e8e8f0">{missing_percent}%</b></span>
                <span>⧉ Duplicates: <b style="color:#e8e8f0">{duplicate_percent}%</b></span>
                <span style="margin-left:auto;color:{health_color};font-weight:600">{health_label}</span>
            </div>
        </div>""", unsafe_allow_html=True)

        # ── Preview: Step 4 ──────────────────────────────────────────────
        section_header("📊","Data Analysis")
        st.markdown('<div class="pill">Statistical Summary</div>', unsafe_allow_html=True)
        st.dataframe(data.describe().round(2), use_container_width=True)

        st.write(f"**Select column for analysis:** {column}")

        st.markdown('<div class="pill" style="margin-top:0.8rem">Value Counts</div>', unsafe_allow_html=True)
        vc_df2 = data[column].value_counts().reset_index(); vc_df2.columns=[column,"Count"]; vc_df2.index=range(1,len(vc_df2)+1)
        st.dataframe(vc_df2, use_container_width=True)

        # ── Preview: Step 5 ──────────────────────────────────────────────
        section_header("📈","Visualization")
        box("Graph shows actual counts of each value — exact data, no ranges.","info")
        chart_data_prev = data[selected_col].value_counts().reset_index(); chart_data_prev.columns=[selected_col,"Count"]
        render_chart(chart_data_prev, data, selected_col, chart_type, key_suffix="_preview")
        st.markdown('<div class="pill" style="margin-top:0.5rem">📋 Graph Data Table</div>', unsafe_allow_html=True)
        st.caption("Exact count of each value in the selected column.")
        st.dataframe(chart_data_display, use_container_width=True)

        # ── Preview: Step 6 ──────────────────────────────────────────────
        section_header("💡","Advanced Recommendations")
        if len(value_counts) > 0:
            box(f"<b>{value_counts.idxmax()}</b> is the highest performing value in '{selected_col}'.","success")
            if len(value_counts)>1: box(f"<b>{value_counts.idxmin()}</b> is the lowest performing value — improvement needed.","warning")
            ca2, cb2 = st.columns(2)
            with ca2:
                st.markdown('<div class="pill">🔝 Top 3 Values</div>', unsafe_allow_html=True)
                t3p = value_counts.head(3).reset_index(); t3p.columns=[selected_col,"Count"]; t3p.index=range(1,len(t3p)+1)
                st.dataframe(t3p, use_container_width=True)
            with cb2:
                st.markdown('<div class="pill">🔻 Bottom 3 Values</div>', unsafe_allow_html=True)
                b3p = value_counts.tail(3).reset_index(); b3p.columns=[selected_col,"Count"]; b3p.index=range(1,len(b3p)+1)
                st.dataframe(b3p, use_container_width=True)
            st.markdown('<div class="pill" style="margin-top:1rem">📊 Percentage Contribution</div>', unsafe_allow_html=True)
            st.dataframe(percent_df, use_container_width=True)
            mxp = (value_counts/value_counts.sum()*100).max()
            if mxp>60:    box(f"<b>{value_counts.idxmax()}</b> dominates with {round(mxp,2)}% — reduce dependency for better balance.","error")
            elif mxp>40:  box(f"<b>{value_counts.idxmax()}</b> has strong influence ({round(mxp,2)}%) — maintain but diversify.","warning")
            else:         box("Distribution is fairly balanced across all categories.","success")
            st.markdown(f'<div class="insight-card"><div class="insight-card-title">🧠 Smart Suggestion</div><div class="insight-card-body">💡 {smart_suggestion_text}</div></div>', unsafe_allow_html=True)

        # ── Preview: Step 7 ──────────────────────────────────────────────
        section_header("🧠","Auto Insight Summary")
        if insight_text: st.markdown(f'<div class="insight-card"><div class="insight-card-title">📊 Insight</div><div class="insight-card-body">👉 {insight_text}</div></div>', unsafe_allow_html=True)
        if key_obs_text: st.markdown(f'<div class="insight-card" style="margin-top:0.8rem"><div class="insight-card-title">📌 Key Observations</div><div class="insight-card-body">💡 {key_obs_text}</div></div>', unsafe_allow_html=True)

        # ── Download Buttons ─────────────────────────────────────────────
        section_header("⬇️","Download Your Report")
        word_bytes = generate_word_report(data, selected_col, chart_type, clean_percent, missing_percent,
                                          duplicate_percent, value_counts, percent_df, top_percent,
                                          insight_text, key_obs_text, smart_suggestion_text,
                                          chart_img_bytes, missing_df, column)
        pdf_bytes  = generate_pdf_report(data, selected_col, chart_type, clean_percent, missing_percent,
                                         duplicate_percent, value_counts, percent_df, top_percent,
                                         insight_text, key_obs_text, smart_suggestion_text,
                                         chart_img_bytes, missing_df, column)
        dc1, dc2 = st.columns(2)
        with dc1:
            if word_bytes:
                st.download_button(label="📝 Download Word Report (.docx)", data=word_bytes,
                                   file_name="Analysis_Report.docx",
                                   mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            else:
                box("Word generation failed. Run: pip install python-docx","error")
        with dc2:
            if pdf_bytes:
                st.download_button(label="📄 Download PDF Report (.pdf)", data=pdf_bytes,
                                   file_name="Analysis_Report.pdf", mime="application/pdf")
            else:
                box("PDF generation failed. Run: pip install reportlab","error")

# ── FOOTER ──────────────────────────────────────────────────────────────────
st.markdown("""
<div style="text-align:center;padding:3rem 0 1rem;color:#2e2e4a;font-size:0.78rem;font-family:'DM Sans',sans-serif">
    <div style="width:40px;height:1px;background:rgba(99,110,250,0.2);margin:0 auto 1rem"></div>
    Intelligent Data Analytics System
</div>
""", unsafe_allow_html=True)
